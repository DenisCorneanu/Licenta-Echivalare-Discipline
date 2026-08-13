# backend/export_docx.py
import json
import sys
import unicodedata

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


def norm(value):
    """Cheie de comparatie: lowercase, fara diacritice si punctuatie."""
    if value is None:
        return ""

    text = unicodedata.normalize("NFKD", str(value))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower()
    text = "".join(ch if (ch.isalnum() or ch.isspace()) else " " for ch in text)
    return " ".join(text.split())


def read_payload():
    raw = sys.stdin.buffer.read()
    data = json.loads(raw.decode("utf-8"))

    if len(sys.argv) < 2:
        print("Missing output path", file=sys.stderr)
        sys.exit(1)

    return data, sys.argv[1]


def table_has_program_headers(table):
    """
    Detecteaza un tabel de echivalare pe baza celor doua coloane principale.
    Cautam in primele trei randuri pentru template-uri cu header pe doua randuri.
    """
    if not table.rows:
        return False

    header_rows = table.rows[: min(3, len(table.rows))]
    joined = " | ".join(
        norm(cell.text)
        for row in header_rows
        for cell in row.cells
    )

    return (
        "disciplina promovata" in joined
        and "disciplina echivalata" in joined
    )


def get_header_texts(table):
    """
    Construieste textul de header pentru fiecare coloana, combinand primele
    trei randuri. Este mai robust pentru celule Word unite pe verticala.
    """
    if not table.rows:
        return []

    column_count = max(len(row.cells) for row in table.rows[: min(3, len(table.rows))])
    texts = []

    for column_index in range(column_count):
        parts = []

        for row in table.rows[: min(3, len(table.rows))]:
            if column_index < len(row.cells):
                text = norm(row.cells[column_index].text)
                if text and text not in parts:
                    parts.append(text)

        texts.append(" ".join(parts))

    return texts


def header_map(table):
    """Gaseste coloanele disciplinei, notei si ECTS-ului."""
    headers = get_header_texts(table)
    columns = {}

    for index, header in enumerate(headers):
        if not header:
            continue

        if "disciplina promovata" in header or "conform planului" in header:
            columns.setdefault("left_name", index)

        if "disciplina echivalata" in header:
            columns.setdefault("right_name", index)

    # Fallback pentru unele template-uri in care celulele unite ascund textul.
    if "left_name" not in columns and len(headers) >= 2:
        columns["left_name"] = 1

    if "right_name" not in columns and len(headers) >= 7:
        columns["right_name"] = 6

    if "left_name" not in columns or "right_name" not in columns:
        return {}

    left_name = columns["left_name"]
    right_name = columns["right_name"]

    grade_columns = [
        index
        for index, header in enumerate(headers)
        if "nota" in header or "calificativ" in header
    ]
    ects_columns = [
        index
        for index, header in enumerate(headers)
        if "credite" in header or "ects" in header
    ]

    # In formular, campurile sursa sunt in stanga disciplinei tinta,
    # iar campurile tinta sunt in dreapta.
    left_grades = [index for index in grade_columns if index < right_name]
    right_grades = [index for index in grade_columns if index > right_name]
    left_ects = [index for index in ects_columns if index < right_name]
    right_ects = [index for index in ects_columns if index > right_name]

    if left_grades:
        columns["left_grade"] = max(left_grades)
    if right_grades:
        columns["right_grade"] = min(right_grades)
    if left_ects:
        columns["left_ects"] = max(left_ects)
    if right_ects:
        columns["right_ects"] = min(right_ects)

    return columns


def build_match_index(rows):
    """
    Index global dupa disciplina tinta.

    Numele disciplinei tinta este unic intr-un template (constrangere in DB),
    deci nu este nevoie sa presupunem ca exista exact doi sau trei ani.
    Astfel, toate tabelele gasite in DOCX pot fi populate.
    """
    matches = {}

    for row in rows:
        target_name = norm(row.get("course_name"))

        # Exportam doar potrivirile efectiv selectate/acceptate.
        # Sugestiile "needs_review" nu trebuie introduse automat in decizie.
        student_name = row.get("student_course_name")
        if not target_name or not student_name:
            continue

        matches[target_name] = {
            "student_course_name": student_name,
            "student_grade": row.get("student_grade"),
            "student_ects": row.get("student_ects"),
            "target_ects": row.get("ects"),
        }

    return matches


def set_cell_text(cell, value):
    """Inlocuieste textul unei celule fara a adauga paragrafe suplimentare."""
    text = "" if value is None else str(value)

    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]

    for paragraph_item in cell.paragraphs:
        for run in paragraph_item.runs:
            run.clear()
        paragraph_item.text = ""

    run = paragraph.add_run(text)
    run.font.size = Pt(11)


def fill_program_table(table, matches):
    """
    Pentru fiecare rand de date:
    disciplina tinta din coloana dreapta -> match global -> campurile din stanga.
    """
    columns = header_map(table)
    if "right_name" not in columns or "left_name" not in columns:
        return 0

    right_name_index = columns["right_name"]
    left_name_index = columns["left_name"]
    filled = 0

    # Pornim de la randul 1; daca exista un al doilea header,
    # nu are o disciplina tinta valida si este sarit automat.
    for row in table.rows[1:]:
        cells = row.cells

        if (
            right_name_index >= len(cells)
            or left_name_index >= len(cells)
        ):
            continue

        target_name = norm(cells[right_name_index].text)
        if not target_name:
            continue

        match = matches.get(target_name)
        if not match:
            continue

        set_cell_text(cells[left_name_index], match["student_course_name"])

        if "left_grade" in columns and columns["left_grade"] < len(cells):
            set_cell_text(cells[columns["left_grade"]], match["student_grade"])

        if (
            "left_ects" in columns
            and columns["left_ects"] < len(cells)
            and match["student_ects"] is not None
        ):
            set_cell_text(cells[columns["left_ects"]], match["student_ects"])

        if "right_grade" in columns and columns["right_grade"] < len(cells):
            set_cell_text(cells[columns["right_grade"]], match["student_grade"])

        if (
            "right_ects" in columns
            and columns["right_ects"] < len(cells)
            and match["target_ects"] is not None
        ):
            set_cell_text(cells[columns["right_ects"]], match["target_ects"])

        filled += 1

    return filled


def find_program_tables(document):
    """Returneaza toate tabelele de echivalare, in ordinea documentului."""
    return [
        table
        for table in document.tables
        if table_has_program_headers(table)
    ]


def format_ects_total(value):
    """Formateaza totalul fara zecimale inutile."""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "0"

    if number.is_integer():
        return str(int(number))

    return f"{number:.4f}".rstrip("0").rstrip(".")


def build_course_year_index(rows):
    """Index disciplina tinta normalizata -> anul disciplinei."""
    result = {}

    for row in rows:
        course_name = norm(row.get("course_name"))
        year = row.get("year")

        if not course_name:
            continue

        try:
            year_number = int(year)
        except (TypeError, ValueError):
            continue

        if year_number > 0:
            result[course_name] = year_number

    return result


def detect_program_table_year(table, course_year_index):
    """
    Detecteaza anul unui tabel dupa disciplinele tinta din coloana dreapta.
    Foloseste votul majoritar, fara a depinde de pozitia fixa a tabelului.
    """
    columns = header_map(table)
    right_name_index = columns.get("right_name")

    if right_name_index is None:
        return None

    votes = {}

    for row in table.rows[1:]:
        if right_name_index >= len(row.cells):
            continue

        target_name = norm(row.cells[right_name_index].text)
        year = course_year_index.get(target_name)

        if year is not None:
            votes[year] = votes.get(year, 0) + 1

    if not votes:
        return None

    return max(votes.items(), key=lambda item: item[1])[0]


def iter_document_blocks(document):
    """Parcurge paragrafele si tabelele in ordinea reala din document."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def fill_ects_totals(document, rows, totals_payload):
    """
    Completeaza numai liniile "Total credite ECTS echivalate".

    Totalul pentru un an este calculat in backend din ECTS-ul disciplinei tinta,
    numai pentru match-uri selectate si promovate (nota >= 5 sau P).
    Campul final "Nr. total de credite care urmeaza a fi obtinute" ramane gol.
    """
    totals_by_year = (totals_payload or {}).get("by_year") or {}
    course_year_index = build_course_year_index(rows)
    pending_year = None
    written = {}

    for block in iter_document_blocks(document):
        if isinstance(block, Table) and table_has_program_headers(block):
            pending_year = detect_program_table_year(
                block,
                course_year_index,
            )
            continue

        if not isinstance(block, Paragraph) or pending_year is None:
            continue

        if not norm(block.text).startswith("total credite ects echivalate"):
            continue

        total_value = totals_by_year.get(str(pending_year), 0)
        total_text = format_ects_total(total_value)

        if ":" in block.text:
            prefix = block.text.split(":", 1)[0].rstrip() + ":"
        else:
            prefix = "Total credite ECTS echivalate:"

        set_paragraph_text_preserving_first_run(
            block,
            f"{prefix} {total_text}",
        )

        written[str(pending_year)] = total_value
        pending_year = None

    return written


def set_paragraph_text_preserving_first_run(paragraph, text):
    """
    Inlocuieste textul unui paragraf si pastreaza formatarea primului run.
    """
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return

    paragraph.add_run(text)


def replace_value_after_label(document, normalized_label, value):
    """
    Cauta un paragraf dupa eticheta normalizata si scrie valoarea dupa ':'.
    Returneaza True daca eticheta a fost gasita.
    """
    for paragraph in document.paragraphs:
        current_text = paragraph.text or ""
        current_norm = norm(current_text)

        if not current_norm.startswith(normalized_label):
            continue

        if ":" in current_text:
            prefix = current_text.split(":", 1)[0].rstrip() + ":"
        else:
            prefix = current_text.rstrip()

        set_paragraph_text_preserving_first_run(
            paragraph,
            f"{prefix} {value}".rstrip(),
        )
        return True

    return False


def fill_student_metadata(document, student_data):
    """
    Numele studentului este obligatoriu.

    Programul destinatie este completat automat daca exista in baza de date.
    Promotia ramane optionala si nu este modificata daca nu are valoare.

    Facultatea de provenienta si domeniul raman neschimbate in template.
    """
    student_name = str(student_data.get("student_name") or "").strip()

    if not student_name:
        raise ValueError("Camp obligatoriu lipsa: Student/a")

    if not replace_value_after_label(document, "student a", student_name):
        raise ValueError("Nu am gasit in template eticheta: Student/a")

    optional_fields = [
        (
            "programul de studiu la care se reinmatriculeaza",
            student_data.get("destination_program"),
        ),
        (
            "promotia cu care continua",
            student_data.get("continuation_cohort"),
        ),
    ]

    for normalized_label, value in optional_fields:
        clean_value = str(value or "").strip()

        if clean_value:
            replace_value_after_label(
                document,
                normalized_label,
                clean_value,
            )


def main():
    data, output_path = read_payload()
    template_path = data["template_path"]
    rows = data["rows"]
    student_data = data.get("student_data") or {}
    ects_totals = data.get("ects_totals") or {}

    document = Document(template_path)
    fill_student_metadata(document, student_data)
    matches = build_match_index(rows)
    program_tables = find_program_tables(document)

    filled_total = 0

    # Grija mare sa nU limitez la primele doua tabele
    # Un template poate avea 1, 2, 3 sau mai multe sectiuni/anuri.
    for table in program_tables:
        filled_total += fill_program_table(table, matches)

    written_ects_totals = fill_ects_totals(
        document,
        rows,
        ects_totals,
    )

    sys.stderr.write(
        f"export_docx.py: tables={len(program_tables)}, "
        f"matches_written={filled_total}, "
        f"ects_totals={written_ects_totals}, "
        f"ects_skipped={ects_totals.get('skipped_courses', [])}\n"
    )

    document.save(output_path)


if __name__ == "__main__":
    main()
